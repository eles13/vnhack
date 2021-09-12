import os
import time
import tqdm
import pandas as pd
import joblib
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
from scipy.sparse import hstack, vstack, csc_matrix
from pymorphy2 import MorphAnalyzer
morph = MorphAnalyzer()
import nltk
nltk.download('stopwords')
from nltk.corpus import stopwords
stopwd = stopwords.words('russian')
import numpy as np
import json

from sklearn.metrics import f1_score
from sklearn.model_selection import train_test_split

from lightautoml.automl.presets.tabular_presets import TabularAutoML, TabularUtilizedAutoML
from lightautoml.tasks import Task
from lightautoml.automl.presets.text_presets import TabularNLPAutoML

RANDOM_STATE = 42
np.random.seed(RANDOM_STATE)

models_path = 'dataset/'
TFIDF_PATH = models_path + 'tfidf.m'
CLUSTERS_PATH = models_path + 'clusters.npy'
AUTOML_PATH = models_path + 'automl.pkl'
AUTOML_NLP_BINARY = models_path + 'binary_classifiсation.pkl'
DOCUMENTS_PATH = './documents/'
DATA_PATH = 'dataset/all_data.csv'
DATASET_PATH = 'dataset/DataSet_Razmetra/'

with open('dataset/corrupt_new.json') as fin:
    alldata = json.load(fin)


NUM_CLASSES = 11
COLORS = [WD_COLOR_INDEX.YELLOW, WD_COLOR_INDEX.BLUE, WD_COLOR_INDEX.BRIGHT_GREEN, WD_COLOR_INDEX.DARK_BLUE,
         WD_COLOR_INDEX.DARK_RED, WD_COLOR_INDEX.DARK_YELLOW, WD_COLOR_INDEX.GRAY_25, WD_COLOR_INDEX.GREEN,
         WD_COLOR_INDEX.PINK, WD_COLOR_INDEX.RED, WD_COLOR_INDEX.VIOLET]

colors_map = {'YELLOW':'3_1', 'BLUE':'3_7', 'BRIGHT_GREEN': '4_1', 'DARK_BLUE': '3_9', 'DARK_RED': '3_5',
              'DARK_YELLOW': '3_3', 'GRAY_25': '3_2', 'GREEN': '4_2', 'PINK':  '4_3', 'RED': '3_6', 'VIOLET': '3_4'}
classes =['3_1', '3_7', '4_1', '3_9', '3_5', '3_3', '3_2', '4_2', '4_3', '3_6', '3_4']


ACTIVATION_CUTOFF = 0.5
DOUBLE_CUTOFF = 0.015

clst = np.load(CLUSTERS_PATH, allow_pickle=True)
tfidf = joblib.load(TFIDF_PATH)
automl = joblib.load(AUTOML_PATH)
morph = MorphAnalyzer()
binary = joblib.load(AUTOML_NLP_BINARY)


vectors = {x[0]: x[1] for x in clst}

os.makedirs(DOCUMENTS_PATH, exist_ok=True)

N_THREADS = 4
N_FOLDS = 3
RANDOM_STATE = 42
TEST_SIZE = 0.05
TIMEOUT = 2 * 3600
TARGET_NAME = 'labels'

def process_document(path, region):
    doc = Document(path)
    data = []
    text=[]
    for par in doc.paragraphs:
        processed = [morph.parse(x)[0].normal_form for x in par.text.split()]
        text.append(' '.join([pr for pr in processed if pr not in stopwd and len(pr) > 1]).lower())
    text_data = pd.DataFrame(text, columns=['text'])
    #binary_mask = binary.predict(text_data).data[:,0]>ACTIVATION_CUTOFF/4
    X = tfidf.transform(text)
    for v in vectors:
        tmp = X.dot(vectors[v].T).mean(axis=0)
        data = np.hstack((data,tmp))
    data = pd.DataFrame([data]).reset_index(drop=True)
    print(path)
    data['regions'] = region
    data.columns = [str(x) for x in data.columns]
    preds = automl.predict(data).data[:][0]
    idx_add_data = []
    labels_toadd = []
    print(path)
    for i,pred in enumerate(preds):
        if pred > ACTIVATION_CUTOFF:
            idx_add_data.append(i)
    maskes=[]
    for v in idx_add_data:
        mask = X.dot(vectors[classes[v]].T).mean(axis=1) > DOUBLE_CUTOFF
        for i in range(len(doc.paragraphs)):
            for j in range(len(doc.paragraphs[i].runs)):
                if mask[i]:# and binary_mask[i]:
                    doc.paragraphs[i].runs[j].font.highlight_color = COLORS[v]
    file_name = path.split('/')[-1]
    doc.save(DOCUMENTS_PATH + file_name.split('.')[0] + '_processed.' + file_name.split('.')[1])
    return True


def train():
    texts = []
    for i in alldata:
        for j in (alldata[i]):
            texts.append(j['text'].lower())
    
    tfidf.fit(texts)
    
    vectors = {}
    for i in alldata:
        kmeans = KMeans(n_clusters=4, random_state=43)
        vector = []
        for j in range(len(alldata[i])):
            alldata[i][j]['emb'] = vectorizer.transform(
                    [alldata[i][j]['text'].lower()])[0]
            vector = vstack((vector, alldata[i][j]['emb']))
        kmeans.fit(vector)
        vectors[i] = kmeans.cluster_centers_
        
    files=[]
    for i in os.walk(DATASET_PATH):
        if len(i[2])>2 and i[2][0].startswith('Edition'):
            files.append(os.path.join(i[0],i[2][0]))
    
    label = {}
    for i,v in enumerate(vectors):
        label[v]=i
    
    with open('corrupt_normed.json') as fin:
        bin_data = json.load(fin)
    bin_texts = []
    bin_labels = []
    for i in bin_data:
        for j in (bin_data[i]):
            if j['type']=='Edition':
                bin_texts.append(j['text'].lower())
                bin_labels.append(1)
            elif j['type']=='NC_Edition':
                bin_texts.append(j['text'].lower())
                bin_labels.append(0)
    data = []
    for file in tqdm.tqdm(files):
        try:
            doc = Document(file)
            text=[]
            for par in doc.paragraphs:
                processed = [morph.parse(x)[0].normal_form for x in par.text.split()]
                text.append(' '.join([pr for pr in processed if pr not in stopwd and len(pr) > 1]).lower())
                bin_texts.append(text[-1])
                bin_labels.append(0)
            X = vectorizer.transform(text)
            s = []
            for v in vectors:
                tmp = X.dot(vectors[v].T).mean(axis=0)
                s = np.hstack((s,tmp))
            data.append(s)
        except:
            print(file)
            continue
    
    bin_data = pd.DataFrame(bin_texts, columns=['text'])
    bin_data['labels'] = bin_labels
    task = Task('binary', )
    roles = {'target': TARGET_NAME,
         'text': ['text']
         }
    automl = TabularNLPAutoML(task = task, 
                           timeout = TIMEOUT,
                           cpu_limit = N_THREADS,
                           general_params = {'use_algos': [['lgb', 'lgb_tuned', 'cb']]},
                           reader_params = {'cv': N_FOLDS, 'n_jobs': N_THREADS, 'random_state': RANDOM_STATE},
                           linear_pipeline_params = {'text_features': "tfidf"},
                           gbm_pipeline_params = {'text_features': "tfidf"},
                           text_params = {'lang': 'ru'}
                          )
    automl.fit_predict(bin_data, roles = roles)
    joblib.dump(automl, AUTOML_NLP_BINARY)
    labels = []
    idx = []
    regs = []
    for file in tqdm.tqdm(files):
        try:
            file = file.split('/')
            labels.append(file[6])
            regs.append(files[5])
            idx.append(file[7])
        except Exception as e:
            print(e, file)
            continue
    data_docx = pd.DataFrame(data)
    data_docx.index = idx
    data_docx['labels'] = labels
    data_docx['regions'] = regs
            
    ndata = []
    for ind in parind:
        processed = [morph.parse(x)[0].normal_form for x in doc.paragraphs[ind].text.split()]
        text = ' '.join([pr for pr in processed if pr not in stopwd and len(pr) > 1]).lower()
        vec = vectorizer.transform(text)
        s = []
        for v in vectors:
            tmp = vec.dot(vectors[v].T).mean(axis=0)
            s = np.hstack((s,tmp))
        ndata.append(s)
    ndata = pd.DataFrame(ndata)
    ndata['labels'] = labels_ind
    ndata['regions'] = region
    
    data = data_docx.append(ndata)
    
    data = data[~data.labels.isin(['3_8', '0_0'])]
    
    tr_data, te_data = train_test_split(data, 
                                    test_size=TEST_SIZE, 
                                    stratify=data[TARGET_NAME], 
                                    random_state=RANDOM_STATE)
    task = Task('multiclass',metric='crossentropy' )
    roles = {'target': TARGET_NAME,
         'drop': [],
         'category': ['regions']
         }
    
    automl = TabularAutoML(task=task,
                          memory_limit=1,
                          timeout=TIMEOUT,
                          cpu_limit=N_THREADS,
                          general_params={'use_algos':
                                              [['lgb', 'cb','lgb_tuned'],['linear_l2']
                                               ]},
                          reader_params={'cv': N_FOLDS,
                                         'random_state': RANDOM_STATE,
                                         'n_jobs': N_THREADS},
                          )
    oof_pred = automl.fit_predict(tr_data, roles=roles)
    
    joblib.dump(vectorizer, TFIDF_PATH)
    joblib.dump(automl, AUTOML_PATH)
    np.save(CLUSTERS_PATH, np.array([np.array([key, vectors[key]]) for key in vectors.keys()]))
    return True

def finetune_on_feedback(path, feedback, region):
    global vectors
    global alldata
    gdoc = Document(path)
    cctr = 0
    parind = []
    labels_ind = []
    for i, par in enumerate(gdoc.paragraphs):
        if par.runs[0].font.highlight_color.split()[0] in colors_map:
            if feedback[cctr]:
                processed = [morph.parse(x)[0].normal_form for x in par.text.split()]
                text = ' '.join([pr for pr in processed if pr not in stopwd and len(pr) > 1]).lower()
                alldata[colors_map[par.runs[0].font.highlight_color.split()[0]]].append({'id': 
                                                                                         len(os.listdir(DOCUMENTS_PATH)) + np.random.randint(0,10000),
                                                                                        'reg': region, 'text': text})
                parind.append(i)
                labels_ind.append(colors_map[par.runs[0].font.highlight_color.split()[0]])
            cctr += 1
    train()
    return True
